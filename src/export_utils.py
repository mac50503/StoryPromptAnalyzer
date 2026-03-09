"""Utilidades para exportar análisis a diferentes formatos."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import re
from datetime import datetime


class ExportManager:
    """Gestor de exportación de análisis."""
    
    @staticmethod
    def export_to_docx(content: str, filename: str, story_id: str = None):
        """
        Exporta el contenido a formato DOCX.
        
        Args:
            content: Contenido del análisis
            filename: Ruta del archivo de salida
            story_id: ID de la historia (opcional)
        """
        doc = Document()
        
        # Título principal
        title = doc.add_heading('Story Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        if story_id:
            meta = doc.add_paragraph()
            meta.add_run(f'Story ID: {story_id}\n').bold = True
            meta.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espacio
        
        # Procesar contenido línea por línea
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            
            # Separadores
            elif line.startswith('===') or line.startswith('---'):
                doc.add_paragraph('_' * 80)
            
            # Listas
            elif re.match(r'^\d+\.\s', line):
                doc.add_paragraph(line, style='List Number')
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            
            # USER/AI labels
            elif line.startswith('USER:'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.color.rgb = RGBColor(41, 128, 185)
            elif line.startswith('AI:') or line.startswith('ASSISTANT:'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.color.rgb = RGBColor(39, 174, 96)
            
            # Texto normal
            else:
                # Procesar bold
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
        
        doc.save(filename)
    
    @staticmethod
    def export_to_pdf(content: str, filename: str, story_id: str = None):
        """
        Exporta el contenido a formato PDF.
        
        Args:
            content: Contenido del análisis
            filename: Ruta del archivo de salida
            story_id: ID de la historia (opcional)
        """
        doc = SimpleDocTemplate(filename, pagesize=letter,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)
        
        # Estilos
        styles = getSampleStyleSheet()
        
        # Estilos personalizados
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=RGBColor(26, 26, 26),
            spaceAfter=30,
            alignment=TA_CENTER
        ))
        
        styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=RGBColor(44, 62, 80),
            spaceAfter=12,
            spaceBefore=12
        ))
        
        styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=RGBColor(52, 73, 94),
            spaceAfter=10,
            spaceBefore=10
        ))
        
        styles.add(ParagraphStyle(
            name='Metadata',
            parent=styles['Normal'],
            fontSize=10,
            textColor=RGBColor(127, 140, 141),
            alignment=TA_CENTER,
            spaceAfter=20
        ))
        
        styles.add(ParagraphStyle(
            name='UserLabel',
            parent=styles['Normal'],
            fontSize=11,
            textColor=RGBColor(41, 128, 185),
            fontName='Helvetica-Bold',
            spaceAfter=6
        ))
        
        styles.add(ParagraphStyle(
            name='AILabel',
            parent=styles['Normal'],
            fontSize=11,
            textColor=RGBColor(39, 174, 96),
            fontName='Helvetica-Bold',
            spaceAfter=6
        ))
        
        # Contenido
        story = []
        
        # Título
        story.append(Paragraph('Story Analysis Report', styles['CustomTitle']))
        
        # Metadata
        if story_id:
            meta_text = f"<b>Story ID:</b> {story_id}<br/>"
            meta_text += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            story.append(Paragraph(meta_text, styles['Metadata']))
        
        story.append(Spacer(1, 0.2*inch))
        
        # Procesar contenido
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            
            if not line:
                story.append(Spacer(1, 0.1*inch))
                continue
            
            # Escapar caracteres especiales para ReportLab
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Headers
            if line.startswith('# '):
                story.append(Paragraph(line[2:], styles['CustomHeading1']))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], styles['CustomHeading2']))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            
            # Separadores
            elif line.startswith('===') or line.startswith('---'):
                story.append(Spacer(1, 0.05*inch))
                story.append(Paragraph('_' * 80, styles['Normal']))
                story.append(Spacer(1, 0.05*inch))
            
            # USER/AI labels
            elif line.startswith('USER:'):
                story.append(Paragraph(line, styles['UserLabel']))
            elif line.startswith('AI:') or line.startswith('ASSISTANT:'):
                story.append(Paragraph(line, styles['AILabel']))
            
            # Listas y texto normal
            else:
                # Procesar bold
                line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                story.append(Paragraph(line, styles['Normal']))
        
        # Generar PDF
        doc.build(story)
